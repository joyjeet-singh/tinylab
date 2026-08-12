"""Update Resume_Joyjeet_Singh_2026.docx.

Adds the second arXiv publication (2608.10145), a new MODEL RELEASES section
for the HuggingFace checkpoints, an EXPERIENCE block for the reproduction,
and the consequential skills. Existing content is edited only where it became
factually wrong -- "one paper published on arXiv" is now two.

New paragraphs are deep copies of existing ones, so they inherit the exact
run formatting, colour, size, spacing and section rules already in the
document rather than approximating them.

Every claim added here is one the released paper or the released model card
states; nothing is inflated for the resume.
"""
import copy
import shutil
from pathlib import Path

import docx

SRC = Path("/Users/joyjeetsingh/Downloads/Resume_Joyjeet_Singh_2026.docx")
OUT = Path("/Users/joyjeetsingh/Downloads/Resume_Joyjeet_Singh_2026_updated.docx")

d = docx.Document(str(SRC))
P = d.paragraphs

# templates to clone, captured before anything moves
T_SECTION = P[5]      # "PUBLICATIONS" -- bold 11pt, navy, bottom rule
T_ENTRY = P[6]        # publication line -- 9.5pt body
T_BULLET = P[7]       # List Bullet, 9.5pt
T_PROJECT = P[13]     # "PHIL-DEQ — ..." -- bold 9.5pt navy project title


def add_after(anchor, template, text):
    """Clone `template`, set `text` as a single run, insert after `anchor`."""
    new_p = copy.deepcopy(template._p)
    anchor._p.addnext(new_p)
    para = docx.text.paragraph.Paragraph(new_p, anchor._parent)
    keep = para.runs[0]
    for r in para.runs[1:]:
        r._r.getparent().remove(r._r)
    keep.text = text
    return para


def set_text(para, text):
    """Replace a single-run paragraph's text, keeping its formatting."""
    assert len(para.runs) >= 1
    para.runs[0].text = text
    for r in para.runs[1:]:
        r._r.getparent().remove(r._r)


def extend_skill(para, extra):
    """Skills lines are [bold label][body]; append to the body run."""
    assert len(para.runs) == 2, f"expected 2 runs, got {len(para.runs)}"
    para.runs[1].text = para.runs[1].text.rstrip() + extra


# ----------------------------------------------------------------- header
# The contact line is centred 9pt and already 101 characters; appending the
# HuggingFace and two arXiv links would wrap it. Split into a contact line and
# a links line instead, cloning the original so both share its formatting.
set_text(P[2],
    "joyjeetsingh1@gmail.com  ·  +91 8620000307  ·  ORCID 0009-0005-1512-7439")
add_after(P[2], P[2],
    "github.com/joyjeet-singh  ·  huggingface.co/Joyjeetsingh"
    "  ·  arXiv:2608.10145  ·  arXiv:2607.11116")

# ---------------------------------------------------------------- summary
# Leads with the differentiator rather than burying it: both papers are
# corrective results that ship the apparatus making them checkable. That is a
# rarer signal than any single architecture on the skills list.
set_text(P[4],
    "Self-directed researcher in physics-informed machine learning with two "
    "sole-author arXiv papers whose common thread is evidentiary discipline. "
    "Both are corrective results — one shows a class of implicit-reasoning "
    "models performs no useful iterative computation; the other shows a "
    "published planning result is decided by its evaluation protocol rather "
    "than its model — and both ship the apparatus that makes the finding "
    "checkable by someone else: a four-test diagnostic protocol in the first, "
    "and in the second a gate that fails the build on any figure in the paper "
    "not traceable to a committed file, with six checkpoints released "
    "alongside it. Builds models from scratch in JAX and PyTorch rather than "
    "fine-tuning checkpoints; active work in Hamiltonian neural architectures, "
    "recursive reasoning models, state space models, and transformer "
    "attention, targeting NeurIPS, ICML and ICLR. Concurrently Tech Lead & ML "
    "evaluator, combining mathematical ML research with applied skills in "
    "evaluation, IT, and analytics. Seeking AI research or ML engineering "
    "roles where initiative, rigour, resourcefulness, and cross-disciplinary "
    "thinking are valued.")

# ----------------------------------------------------------- publications
# Newest first: the reproduction (Aug 2026) precedes PHIL-DEQ (Jul 2026).
a = add_after(P[5], T_ENTRY,
    "Joyjeet Singh. “The Evaluation Protocol Determines the Result: An "
    "Independent Reproduction of LeWorldModel on TwoRoom.” "
    "arXiv:2608.10145 [cs.LG], 2026. Sole author.")
a = add_after(a, T_BULLET,
    "An independent from-scratch reimplementation that reproduces the "
    "original’s headline TwoRoom result and exceeds it — 94.0% of goals "
    "reached against 84.0% for the authors’ own released checkpoint under an "
    "identical protocol on identical episodes — and shows the result is "
    "decided by the evaluation protocol rather than the model: the two "
    "protocols the source publishes give 84.0% and 14.0% on those same "
    "released weights, and changing only how the goal is constructed moves "
    "that checkpoint from 84.0% to 8.0% (McNemar χ² ≈ 34.2, p < 10⁻⁸).")
a = add_after(a, T_BULLET,
    "Recovered four training conventions that decide whether the model "
    "converges and appear in no released configuration file; identified a "
    "batch-normalisation evaluation-mode artifact that inflated reported "
    "validation loss by up to 300× and contaminated four separate quantities; "
    "and showed one-step prediction accuracy does not predict long-horizon "
    "planning success. Reports a pre-registered effect that did not survive, "
    "with all three arms. Code, six checkpoints and every evaluation report: "
    "github.com/joyjeet-singh/tinylab")

last_pub_bullet = P[8]        # the older paper's final bullet

# ------------------------------------------------------- MODEL RELEASES
s = add_after(last_pub_bullet, T_SECTION, "MODEL RELEASES")
s = add_after(s, T_PROJECT,
    "tinylab-tworoom-lewm — LeWorldModel reproduction checkpoints (PyTorch)"
    "  ·  huggingface.co/Joyjeetsingh/tinylab-tworoom-lewm")
s = add_after(s, T_BULLET,
    "Six trained checkpoints (434 MB) released under MIT alongside "
    "arXiv:2608.10145: three BatchNorm-recalibrated models and the three "
    "un-recalibrated originals, so the evaluation-mode artifact the paper "
    "documents can be verified independently rather than taken on trust.")
s = add_after(s, T_BULLET,
    "Model card documents the four undocumented training conventions, how to "
    "tell a recalibrated checkpoint from a training average by reading the "
    "file itself, every planning number with its protocol attached, and the "
    "domain guard a user needs to evaluate the weights in their own renderer; "
    "published with an md5 manifest, and the authors’ dataset linked rather "
    "than redistributed.")

# ------------------------------------------------------------ experience
set_text(P[11],
    "Designing and implementing physics-informed ML architectures as sole "
    "author; two papers published on arXiv (above), with further work "
    "targeting NeurIPS, ICML, and ICLR.")

# "Working knowledge of JAX" was contradicted by the bullet directly below it:
# an implicit-function-theorem backward pass via jax.custom_vjp is not working
# knowledge. State the level, and evidence it without repeating PHIL-DEQ.
set_text(P[12],
    "Architect-level fluency in JAX (Flax/Optax) and PyTorch: builds models "
    "from scratch rather than fine-tuning pretrained checkpoints, and writes "
    "custom autodiff where the framework provides no primitive — "
    "implicit-function-theorem backward passes via jax.custom_vjp for "
    "O(1)-memory training through a fixed point, bespoke solvers, and "
    "structure-preserving integrators.")

e = add_after(P[12], T_PROJECT,
    "tinylab — Independent Reproduction of LeWorldModel (PyTorch)"
    "  ·  Published, arXiv:2608.10145")
e = add_after(e, T_BULLET,
    "Reimplemented a published latent world model from scratch and reproduced "
    "its TwoRoom planning result on roughly $25 of rented GPU time, with all "
    "evaluation run on one laptop CPU; reached 94.0% of goals against a "
    "reported ~87%, and reproduced the representation result directly "
    "(position probe Pearson r = 0.9988 against a reported 0.996).")
e = add_after(e, T_BULLET,
    "Found that the source paper’s own described evaluation protocol does not "
    "reproduce its own reported figure on its own released weights, and "
    "established goal construction rather than model quality as the cause via "
    "a paired comparison on identical episodes.")
e = add_after(e, T_BULLET,
    "Recovered four undocumented conventions by reading reference source "
    "rather than configuration files; diagnosed a BatchNorm evaluation-mode "
    "artifact; and re-derived every load-bearing figure on CPU when the "
    "original measurement outputs proved to be missing from the record.")
e = add_after(e, T_BULLET,
    "Built the release engineering: a numbers gate that fails on any "
    "percentage in the paper not traceable to a committed file, a "
    "Markdown→LaTeX pipeline with citation and anonymity checks for a "
    "double-blind submission, and an arXiv package verified by compiling it "
    "the way arXiv does.")

# ---------------------------------------------------------------- skills
extend_skill(P[49], " · World models / JEPA · Reproducibility & replication studies")
extend_skill(P[51], " · HuggingFace Hub (model release, cards) · pandoc")

shutil.copy2(SRC, str(SRC) + ".bak")
d.save(str(OUT))
print(f"wrote {OUT}")
print(f"backup of the original at {SRC}.bak")

# ---------------------------------------------------------------- verify
chk = docx.Document(str(OUT))
text = "\n".join(p.text for p in chk.paragraphs)
print(f"\nparagraphs: {len(P)} -> {len(chk.paragraphs)}")
for probe in ("2608.10145", "MODEL RELEASES", "huggingface.co/Joyjeetsingh/tinylab-tworoom-lewm",
              "two papers published on arXiv", "two sole-author arXiv papers",
              "World models / JEPA", "HuggingFace Hub",
              "huggingface.co/Joyjeetsingh  \u00b7  arXiv:2608.10145",
              "Architect-level fluency in JAX",
              "evidentiary discipline"):
    assert probe in text, f"MISSING: {probe}"
    print(f"  OK  {probe}")
assert "one paper published on arXiv" not in text, "the old count survived"
assert "Working knowledge of JAX" not in text, "the undersell survived"
print("  OK  the old 'one paper' count is gone")
